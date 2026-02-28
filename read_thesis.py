from docx import Document

doc = Document(r"PhD thesis\PhD Oral Examination Proposal_Ver2.docx")

# Extract all text
full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text)

for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text.strip())
        if any(row_text):
            full_text.append(" | ".join(row_text))

output = "\n".join(full_text)
print(output)
